import docx
from pathlib import Path

def clone_and_fill(source_path: Path, changes_json: dict, output_path: Path) -> None:
    """Clone a DOCX file and apply text replacements globally (paragraphs, tables, headers, footers)."""
    doc = docx.Document(source_path)
    
    replacements = changes_json.get("replacements", [])
    if not replacements:
        raise ValueError(
            "clone_and_fill requires a 'replacements' list in the JSON. "
            "Use create_docx() for new documents, or provide replacement pairs for template filling."
        )

    def replace_in_paragraph(p, old_text, new_text):
        if not p or not old_text or old_text not in p.text:
            return
            
        # Try simple run-level replacement first to preserve formatting
        replaced_in_run = False
        for run in p.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, str(new_text))
                replaced_in_run = True
                
        # If the text spans multiple runs, fallback to paragraph-level replacement (loses formatting)
        if not replaced_in_run and old_text in p.text:
            p.text = p.text.replace(old_text, str(new_text))

    def process_section(container):
        """Process paragraphs and tables in a document, header, or footer."""
        for p in container.paragraphs:
            for rep in replacements:
                replace_in_paragraph(p, rep.get("old"), rep.get("new"))
                    
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_section(cell)

    # 1. Process Main Body
    process_section(doc)

    # 2. Process Headers and Footers
    for section in doc.sections:
        if section.header:
            process_section(section.header)
        if section.footer:
            process_section(section.footer)
        # Handle different header/footer types if they exist
        if hasattr(section, 'first_page_header') and section.first_page_header:
            process_section(section.first_page_header)
        if hasattr(section, 'first_page_footer') and section.first_page_footer:
            process_section(section.first_page_footer)
        if hasattr(section, 'even_page_header') and section.even_page_header:
            process_section(section.even_page_header)
        if hasattr(section, 'even_page_footer') and section.even_page_footer:
            process_section(section.even_page_footer)
                        
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Twips

def create_docx(content_json: dict, output_path: Path) -> None:
    """Create a new DOCX file from structured JSON content."""
    doc = docx.Document()
    
    title = content_json.get("title", "")
    if title:
        doc.add_heading(title, 0)
        
    # Handle New Structured Schema
    sections = content_json.get("sections", [])
    for section in sections:
        stype = section.get("type")
        text = section.get("text", "")
        fmt_config = section.get("format", {})
        
        p = None
        if stype == "heading":
            level = section.get("level", 1)
            p = doc.add_heading(text, level)
        elif stype == "paragraph":
            p = doc.add_paragraph(text)
        elif stype == "bullet_list":
            items = section.get("items", [])
            for item in items:
                p_item = doc.add_paragraph(item, style='List Bullet')
                # Apply formatting to list items if provided
                apply_paragraph_format(p_item, fmt_config)
        elif stype == "numbered_list":
            items = section.get("items", [])
            for item in items:
                p_item = doc.add_paragraph(item, style='List Number')
                apply_paragraph_format(p_item, fmt_config)
        elif stype == "table":
            headers = section.get("headers", [])
            rows = section.get("rows", [])
            if headers or rows:
                num_cols = len(headers) if headers else (len(rows[0]) if rows else 0)
                table = doc.add_table(rows=1 if headers else 0, cols=num_cols)
                table.style = 'Table Grid'
                if headers:
                    hdr_cells = table.rows[0].cells
                    for i, h in enumerate(headers):
                        hdr_cells[i].text = h
                for row_data in rows:
                    row_cells = table.add_row().cells
                    for i, val in enumerate(row_data):
                        if i < len(row_cells):
                            row_cells[i].text = str(val)
        
        if p and fmt_config:
            apply_paragraph_format(p, fmt_config)

    # Backward compatibility for old flat format
    content = content_json.get("content", "")
    if content and not sections:
        # Split by newlines and create paragraphs
        paragraphs = content.split('\n')
        for p_text in paragraphs:
            if p_text.strip():
                doc.add_paragraph(p_text.strip())
                
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

def apply_paragraph_format(p, fmt_config: dict):
    """Utility to apply formatting to a paragraph object."""
    if not p or not fmt_config:
        return
        
    pf = p.paragraph_format
    
    if "indent_left" in fmt_config:
        pf.left_indent = Twips(int(fmt_config["indent_left"]))
    
    if "first_line_indent" in fmt_config:
        pf.first_line_indent = Twips(int(fmt_config["first_line_indent"]))
        
    if "alignment" in fmt_config:
        align_str = fmt_config["alignment"].upper()
        if align_str == "LEFT":
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align_str == "CENTER":
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align_str == "RIGHT":
            pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align_str == "JUSTIFY":
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
