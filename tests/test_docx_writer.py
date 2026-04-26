import pytest
import docx
from pathlib import Path

# This will fail to import initially
from tools.docx_writer import edit_docx, create_docx

def test_create_docx(tmp_path):
    output_path = tmp_path / "new.docx"
    content = {
        "title": "Test Title",
        "content": "This is a paragraph.\n\nAnother paragraph."
    }
    create_docx(content, output_path)
    
    assert output_path.exists()
    doc = docx.Document(output_path)
    assert len(doc.paragraphs) >= 2
    assert "Test Title" in doc.paragraphs[0].text
    
def test_edit_docx(tmp_path):
    # Create a source doc first
    source_path = tmp_path / "source.docx"
    doc = docx.Document()
    doc.add_paragraph("Old text here.")
    doc.add_paragraph("Do not touch this.")
    doc.save(source_path)
    
    output_path = tmp_path / "edited.docx"
    changes = {
        "replacements": [
            {"old": "Old text here.", "new": "New text here."}
        ]
    }
    
    edit_docx(source_path, changes, output_path)
    
    assert output_path.exists()
    edited_doc = docx.Document(output_path)
    assert edited_doc.paragraphs[0].text == "New text here."
    assert edited_doc.paragraphs[1].text == "Do not touch this."
